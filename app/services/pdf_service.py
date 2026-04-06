import os
import logging
import fitz  # PyMuPDF
from PIL import Image
from fpdf import FPDF
from docx import Document
from app.services.ocr_service import OCRService
from app.services.translation_service import TranslationService
from config.settings import Config

logger = logging.getLogger(__name__)

class PDFService:
    def __init__(self, gemini_api_key):
        self.translation_service = TranslationService(api_key=gemini_api_key)
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Create necessary directories if they don't exist"""
        directories = [
            Config.RAW_FILES_DIR,
            Config.OUTPUT_FILES_DIR,
            Config.OUTPUT_OCR_FILES_DIR,
            Config.OUTPUT_DOCX_FILES_DIR
        ]
        for directory in directories:
            if not os.path.exists(directory):
                os.makedirs(directory)
    
    def process_pdf(self, input_pdf_path, pdf_name, input_language='en'):
        """Process a PDF: OCR each page, translate, and save results."""
        try:
            logger.info(f"Processing PDF: {pdf_name}")

            # Initialize OCR engine for specified language
            ocr_service = OCRService(lang=input_language)

            # Initialize output documents
            pdf_output = FPDF()
            pdf_output.set_auto_page_break(auto=True, margin=15)
            docx_ocr = Document()
            docx_translated = Document()

            # Configure font for PDF
            font_path = os.path.join(Config.FONT_DIR, "DejaVuSansCondensed.ttf")
            windows_nirmala = "C:\\Windows\\Fonts\\Nirmala.ttf"
            windows_arial = "C:\\Windows\\Fonts\\arial.ttf"
            
            if os.path.exists(font_path):
                pdf_output.add_font('Unicode', '', font_path)
                pdf_output.set_font('Unicode', '', 12)
            elif os.path.exists(windows_nirmala):
                # Nirmala UI natively supports Devanagari, Gujarati, Rupee symbols, etc.
                pdf_output.add_font('Unicode', '', windows_nirmala)
                pdf_output.set_font('Unicode', '', 12)
            elif os.path.exists(windows_arial):
                pdf_output.add_font('Unicode', '', windows_arial)
                pdf_output.set_font('Unicode', '', 12)
            else:
                pdf_output.set_font('helvetica', '', 12)

            # Open PDF — use try/finally so doc.close() ALWAYS runs on Windows
            # (Windows keeps an exclusive lock on open file handles)
            doc = fitz.open(input_pdf_path)
            try:
                total_pages = len(doc)
                for i in range(total_pages):
                    logger.info(f"Processing page {i+1}/{total_pages}")
                    page = doc.load_page(i)

                    # Render to image at 2x resolution for better OCR accuracy
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    # Explicitly free the pixmap memory
                    pix = None

                    ocr_text = ocr_service.extract_text_from_pil(img)
                    translated_text = self.translation_service.translate(ocr_text)

                    # OCR DOCX
                    docx_ocr.add_heading(f"Page {i+1}", level=2)
                    docx_ocr.add_paragraph(ocr_text)
                    docx_ocr.add_page_break()

                    # Translated PDF
                    pdf_output.add_page()
                    pdf_output.cell(200, 10, txt=f"Page {i+1}:", ln=True, align='L')
                    try:
                        pdf_output.multi_cell(0, 7, txt=translated_text)
                    except Exception as fpdf_encoding_err:
                        logger.warning(f"PDF Font Encoding issue on page {i+1}, applying safe fallback: {fpdf_encoding_err}")
                        safe_fallback = translated_text.encode('ascii', 'replace').decode('ascii')
                        pdf_output.multi_cell(0, 7, txt=safe_fallback)

                    # Translated DOCX
                    docx_translated.add_heading(f"Page {i+1}", level=2)
                    docx_translated.add_paragraph(translated_text)
                    docx_translated.add_page_break()
            finally:
                # ALWAYS close the fitz document to release the Windows file lock
                doc.close()

            # Save outputs (only reached if no exception above)
            pdf_output.output(os.path.join(Config.OUTPUT_FILES_DIR, f"{pdf_name}.pdf"))
            docx_ocr.save(os.path.join(Config.OUTPUT_OCR_FILES_DIR, f"{pdf_name}.docx"))
            docx_translated.save(os.path.join(Config.OUTPUT_DOCX_FILES_DIR, f"{pdf_name}.docx"))

            logger.info(f"Successfully processed PDF: {pdf_name}")
            return True

        except Exception as e:
            # Safely encode the error to avoid charmap codec exceptions in Windows console
            safe_err = str(e).encode("ascii", "replace").decode("ascii")
            logger.error(f"Error processing PDF {pdf_name}: {safe_err}")
            return False
