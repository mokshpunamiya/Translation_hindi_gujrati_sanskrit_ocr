import os
import logging
import time
from pathlib import Path
import fitz
from fpdf import FPDF
from docx import Document
import pytesseract
from PIL import Image
from openai import OpenAI
from typing import Optional, Dict, Any
import hashlib
from config import Config

logger = logging.getLogger(__name__)

class PDFProcessor:
    def __init__(self):
        self.client = OpenAI(api_key=Config.OPENAI_API_KEY)
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Create necessary directories if they don't exist"""
        directories = [
            Config.RAW_FILES_DIR,
            Config.OUTPUT_FILES_DIR,
            Config.OUTPUT_OCR_FILES_DIR,
            Config.OUTPUT_DOCX_FILES_DIR
        ]
        for directory in directories:
            Path(directory).mkdir(parents=True, exist_ok=True)
    
    def _get_language_config(self, language: str) -> str:
        """Get Tesseract language configuration"""
        configs = {
            'gujarati': r'--oem 1 --psm 3 -l guj+san',
            'hindi': r'--oem 1 --psm 3 -l hin+san',
            'english': r'--oem 1 --psm 3 -l eng'
        }
        return configs.get(language.lower(), configs['english'])
    
    def _translate_text(self, text: str) -> str:
        """Translate text using OpenAI with retry logic"""
        max_retries = 3
        for attempt in range(max_retries):
            try:
                completion = self.client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "user", "content": f"{text}\nTranslate above text to English. If corrections or clarifications are needed for accuracy just do it, and provide the revised text with appropriate formatting. Also modify text if it doesn't make sense. Ensure your response contains only the translated text and nothing else."}
                    ],
                    timeout=30  # 30 second timeout
                )
                return completion.choices[0].message.content
            except Exception as e:
                logger.error(f"Translation attempt {attempt + 1} failed: {e}")
                if attempt == max_retries - 1:
                    raise
                time.sleep(2 ** attempt)  # Exponential backoff
    
    def _process_page(self, page, page_num: int, language: str) -> Dict[str, Any]:
        """Process a single page of the PDF"""
        temp_image_path = f"temp_page_{page_num}.png"
        try:
            # Extract image from PDF
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution for better OCR
            pix.save(temp_image_path)
            
            # OCR the image
            image = Image.open(temp_image_path)
            custom_config = self._get_language_config(language)
            ocr_text = pytesseract.image_to_string(image, config=custom_config)
            
            # Translate text
            translated_text = self._translate_text(ocr_text)
            
            return {
                'ocr_text': ocr_text,
                'translated_text': translated_text,
                'page_num': page_num
            }
        finally:
            # Clean up temp file
            if os.path.exists(temp_image_path):
                os.remove(temp_image_path)
    
    def process_pdf(self, input_pdf_path: str, pdf_name: str, input_language: str) -> bool:
        """Process PDF with progress tracking"""
        try:
            logger.info(f"Processing PDF: {pdf_name}")
            
            # Initialize documents
            pdf_output = FPDF()
            pdf_output.set_auto_page_break(auto=True, margin=15)
            docx_ocr = Document()
            docx_translated = Document()
            
            # Add fonts if available
            font_path = Path(Config.FONT_DIR) / "DejaVuSansCondensed.ttf"
            if font_path.exists():
                pdf_output.add_font('DejaVu', '', str(font_path), uni=True)
                pdf_output.set_font('DejaVu', '', 12)
            else:
                pdf_output.set_font('helvetica', '', 12)
            
            # Process PDF
            doc = fitz.open(input_pdf_path)
            total_pages = len(doc)
            
            for i, page in enumerate(doc):
                logger.info(f"Processing page {i+1}/{total_pages}")
                result = self._process_page(page, i+1, input_language)
                
                # Add to OCR document
                docx_ocr.add_paragraph(f"Page {i+1}:\n{result['ocr_text']}")
                docx_ocr.add_page_break()
                
                # Add to translated document
                pdf_output.add_page()
                pdf_output.cell(200, 10, txt=f"Page {i+1}:", ln=True, align='L')
                pdf_output.multi_cell(0, 7, txt=result['translated_text'])
                
                docx_translated.add_paragraph(f"Page {i+1}:\n{result['translated_text']}")
                docx_translated.add_page_break()
            
            doc.close()
            
            # Save outputs
            pdf_output.output(os.path.join(Config.OUTPUT_FILES_DIR, f"{pdf_name}.pdf"))
            docx_ocr.save(os.path.join(Config.OUTPUT_OCR_FILES_DIR, f"{pdf_name}.docx"))
            docx_translated.save(os.path.join(Config.OUTPUT_DOCX_FILES_DIR, f"{pdf_name}.docx"))
            
            logger.info(f"Successfully processed PDF: {pdf_name}")
            return True
            
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_name}: {e}")
            return False

processor = PDFProcessor()

def extract_text_and_translate(input_pdf_path: str, pdf_name: str, input_language: str) -> str:
    """Wrapper function for backward compatibility"""
    return 'success' if processor.process_pdf(input_pdf_path, pdf_name, input_language) else 'failed'